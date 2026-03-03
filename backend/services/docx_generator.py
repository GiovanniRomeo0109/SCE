"""
Generatore DOCX per la sicurezza nei cantieri edili italiani.
Conforme al D.Lgs. 81/2008 — Allegati XII e XV.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import Optional
import os


# ═══════════════════════════════════════════════
# PALETTE COLORI
# ═══════════════════════════════════════════════
NAVY   = RGBColor(0x1A, 0x3A, 0x5C)
GOLD   = RGBColor(0xC8, 0x8B, 0x2A)
GREY   = RGBColor(0x5A, 0x6B, 0x7D)
RED    = RGBColor(0xC0, 0x39, 0x2B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BG_H   = "E8F0F8"   # sfondo intestazione tabella (hex)
BG_NAV = "1A3A5C"   # sfondo header (hex)
BG_GLD = "C88B2A"   # sfondo gold (hex)


# ═══════════════════════════════════════════════
# UTILITÀ BASSE
# ═══════════════════════════════════════════════

def _shd(cell, fill_hex: str):
    """Sfondo cella tabella."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def v(d: dict, key: str, default: str = "") -> str:
    """Estrae un valore dal form_data come stringa."""
    val = d.get(key)
    return str(val).strip() if val else default


def _setup(doc: Document):
    """Margini e font base."""
    for s in doc.sections:
        s.left_margin   = Cm(2.5)
        s.right_margin  = Cm(2.5)
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _footer(doc: Document, label: str):
    """Piè di pagina standard."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"{label}  —  Generato il {datetime.now().strftime('%d/%m/%Y %H:%M')}  —  "
        "Conforme D.Lgs. 81/2008  —  Da verificare e completare a cura del professionista"
    )
    run.font.size = Pt(7)
    run.font.color.rgb = GREY
    run.italic = True


# ── Header bicolore ──────────────────────────────────────────────────────────

def _header_bicolor(doc: Document, left_text: str, right_text: str = "D.Lgs. 81/2008"):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell_l = tbl.cell(0, 0)
    _shd(cell_l, BG_NAV)
    p = cell_l.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"  {left_text}")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = WHITE; r.font.name = "Calibri"

    cell_r = tbl.cell(0, 1)
    _shd(cell_r, BG_GLD)
    p2 = cell_r.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(f"{right_text}  ")
    r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = WHITE; r2.font.name = "Calibri"

    doc.add_paragraph()


# ── Titolo sezione ───────────────────────────────────────────────────────────

def _section(doc: Document, title: str, num: str = ""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    txt = f"{num}  {title}".strip()
    r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY; r.font.name = "Calibri"

    # bordo inferiore oro
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "C88B2A")
    bdr.append(bot)
    pPr.append(bdr)
    return p


def _subsection(doc: Document, title: str):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = NAVY; r.font.name = "Calibri"
    return p


def _field(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(value or "—")
    r2.font.size = Pt(10)


def _note(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"Rif. norm.: {text}")
    r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY


def _table(doc: Document, rows_data: list, headers: list = None):
    """Tabella dati formattata."""
    n_cols = len(headers) if headers else (len(rows_data[0]) if rows_data else 2)
    n_rows = len(rows_data) + (1 if headers else 0)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"

    ri = 0
    if headers:
        for ci, h in enumerate(headers):
            cell = tbl.cell(0, ci)
            _shd(cell, BG_H)
            r = cell.paragraphs[0].add_run(h)
            r.bold = True; r.font.size = Pt(9); r.font.color.rgb = NAVY
        ri = 1

    for row in rows_data:
        for ci, val in enumerate(row):
            tbl.cell(ri, ci).paragraphs[0].add_run(str(val) if val else "—").font.size = Pt(9)
        ri += 1

    doc.add_paragraph()
    return tbl


def _warning(doc: Document):
    p = doc.add_paragraph()
    r = p.add_run(
        "⚠  AVVISO: Documento generato automaticamente come supporto al professionista. "
        "Deve essere verificato, completato e firmato dal professionista responsabile "
        "prima dell'utilizzo."
    )
    r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)


def _save(doc: Document, prefix: str, output_dir: str, tag: str = "") -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_tag = tag.replace(" ", "_")[:30] if tag else "cantiere"
    filename  = f"{prefix}_{safe_tag}_{ts}.docx"
    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# NOTIFICA PRELIMINARE
# ═══════════════════════════════════════════════════════════════════════════════

def genera_notifica_preliminare(form_data: dict, contenuto_ai: Optional[dict], output_dir: str) -> str:
    doc = Document()
    _setup(doc)
    _footer(doc, "NOTIFICA PRELIMINARE — Art. 99 D.Lgs. 81/2008 — All. XII")

    # ── Intestazione ──────────────────────────────────────────────────────────
    _header_bicolor(doc, "NOTIFICA PRELIMINARE")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NOTIFICA PRELIMINARE")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("ai sensi dell'Art. 99 — D.Lgs. 9 aprile 2008, n. 81 — Allegato XII")
    r2.italic = True; r2.font.size = Pt(10); r2.font.color.rgb = GREY
    doc.add_paragraph()

    # ── Destinatari ───────────────────────────────────────────────────────────
    _subsection(doc, "DESTINATARI")
    _table(doc,
        [
            ["Azienda Sanitaria Locale (ASL)", v(form_data, "asl_destinataria", "___________________________")],
            ["Ispettorato Territoriale del Lavoro (ITL)", v(form_data, "itl_destinatario", "___________________________")],
        ],
        ["Ente", "Sede territoriale"])

    # ── 1. Dati cantiere ──────────────────────────────────────────────────────
    _section(doc, "DATI DEL CANTIERE", "1.")
    _note(doc, "Allegato XII, punto 2 e 4 — D.Lgs. 81/2008")
    _field(doc, "Indirizzo cantiere", v(form_data, "indirizzo_cantiere"))
    _field(doc, "Comune",   f"{v(form_data,'citta_cantiere')} ({v(form_data,'provincia_cantiere')})")
    _field(doc, "CAP",      v(form_data, "cap_cantiere"))
    _field(doc, "Natura dell'opera", v(form_data, "natura_opera"))
    _field(doc, "Descrizione sommaria", v(form_data, "descrizione_opera"))

    # ── 2. Committente ────────────────────────────────────────────────────────
    _section(doc, "COMMITTENTE", "2.")
    _note(doc, "Allegato XII, punto 3 — D.Lgs. 81/2008")
    if v(form_data, "committente_tipo") == "persona_giuridica":
        _field(doc, "Ragione Sociale", v(form_data, "committente_ragione_sociale"))
        _field(doc, "P.IVA",           v(form_data, "committente_piva"))
        _field(doc, "Legale Rappresentante",
               f"{v(form_data,'committente_nome')} {v(form_data,'committente_cognome')}")
    else:
        _field(doc, "Nome e Cognome",
               f"{v(form_data,'committente_nome')} {v(form_data,'committente_cognome')}")
        _field(doc, "Codice Fiscale", v(form_data, "committente_cf"))
    _field(doc, "Indirizzo", v(form_data, "committente_indirizzo"))
    _field(doc, "Città",
           f"{v(form_data,'committente_citta')} ({v(form_data,'committente_provincia','')}) "
           f"{v(form_data,'committente_cap','')}")
    _field(doc, "Telefono", v(form_data, "committente_telefono"))
    _field(doc, "Email / PEC", v(form_data, "committente_email"))

    # ── 3. Responsabile dei Lavori ────────────────────────────────────────────
    _section(doc, "RESPONSABILE DEI LAVORI", "3.")
    _note(doc, "Allegato XII, punto 5 — D.Lgs. 81/2008")
    if v(form_data, "rl_nome"):
        _field(doc, "Nome e Cognome", f"{v(form_data,'rl_nome')} {v(form_data,'rl_cognome')}")
        _field(doc, "Qualifica",  v(form_data, "rl_qualifica"))
        _field(doc, "Indirizzo",  v(form_data, "rl_indirizzo"))
        _field(doc, "Telefono",   v(form_data, "rl_telefono"))
        _field(doc, "Email/PEC",  v(form_data, "rl_email"))
    else:
        p = doc.add_paragraph("Coincide con il Committente")
        p.runs[0].italic = True

    # ── 4. Coordinatori ───────────────────────────────────────────────────────
    _section(doc, "COORDINATORI PER LA SICUREZZA", "4.")
    _note(doc, "Allegato XII, punti 6 e 7 — D.Lgs. 81/2008")

    _subsection(doc, "4.1  CSP — Coordinatore per la Progettazione")
    if v(form_data, "csp_nome"):
        _field(doc, "Nome e Cognome", f"{v(form_data,'csp_nome')} {v(form_data,'csp_cognome')}")
        _field(doc, "Ordine Professionale", v(form_data, "csp_ordine"))
        _field(doc, "N. Iscrizione",        v(form_data, "csp_numero_ordine"))
        _field(doc, "Telefono",             v(form_data, "csp_telefono"))
        _field(doc, "PEC",                  v(form_data, "csp_pec"))
    else:
        p = doc.add_paragraph("Da designare / Non ancora nominato"); p.runs[0].italic = True

    _subsection(doc, "4.2  CSE — Coordinatore per l'Esecuzione")
    if v(form_data, "cse_nome"):
        _field(doc, "Nome e Cognome", f"{v(form_data,'cse_nome')} {v(form_data,'cse_cognome')}")
        _field(doc, "Ordine Professionale", v(form_data, "cse_ordine"))
        _field(doc, "N. Iscrizione",        v(form_data, "cse_numero_ordine"))
        _field(doc, "Telefono",             v(form_data, "cse_telefono"))
        _field(doc, "PEC",                  v(form_data, "cse_pec"))
    else:
        p = doc.add_paragraph("Da designare prima dell'inizio lavori"); p.runs[0].italic = True

    # ── 5. Dati lavori e lavoratori ───────────────────────────────────────────
    _section(doc, "DATI LAVORI E LAVORATORI", "5.")
    _note(doc, "Allegato XII, punti 8-11 — D.Lgs. 81/2008")
    _field(doc, "Data presunta inizio lavori",    v(form_data, "data_inizio"))
    _field(doc, "Durata presunta lavori",         v(form_data, "durata_lavori"))
    _field(doc, "Durata in uomini-giorno (UG)",   v(form_data, "uomini_giorno"))
    _field(doc, "N. max lavoratori contemporanei", v(form_data, "max_lavoratori"))
    _field(doc, "N. previsto imprese / lav. autonomi", v(form_data, "num_imprese"))

    # ── 6. Imprese selezionate ────────────────────────────────────────────────
    _section(doc, "IMPRESE ESECUTRICI GIÀ SELEZIONATE", "6.")
    _note(doc, "Allegato XII, punto 12 — D.Lgs. 81/2008")
    imprese = form_data.get("imprese_esecutrici", [])
    if imprese:
        _table(doc,
            [[i.get("ragione_sociale",""), i.get("piva",""),
              f"{i.get('nome_dl','')} {i.get('cognome_dl','')}".strip(),
              i.get("attivita","")]
             for i in imprese],
            ["Ragione Sociale", "P.IVA", "Datore di Lavoro", "Attività nel cantiere"])
    else:
        p = doc.add_paragraph("Nessuna impresa ancora selezionata alla data della notifica.")
        p.runs[0].italic = True

    # ── 7. Firme ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    _section(doc, "DICHIARAZIONE E FIRME", "7.")
    p = doc.add_paragraph(
        "Il sottoscritto, nella qualità di Committente / Responsabile dei Lavori, trasmette la "
        "presente Notifica Preliminare ai sensi dell'art. 99 del D.Lgs. 81/2008 e si impegna ad "
        "aggiornarla tempestivamente in caso di variazioni significative."
    )
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    _table(doc, [
        ["Luogo e data", f"{v(form_data,'citta_cantiere')},  _____ / _____ / _______"],
        ["Il Committente / Responsabile dei Lavori", ""],
        ["Firma", "_____________________________"],
    ], ["Voce", "Dettaglio"])

    doc.add_paragraph()
    _warning(doc)

    return _save(doc, "Notifica_Preliminare", output_dir, v(form_data, "citta_cantiere"))


# ═══════════════════════════════════════════════════════════════════════════════
# PSC — PIANO DI SICUREZZA E COORDINAMENTO
# ═══════════════════════════════════════════════════════════════════════════════

def genera_psc(form_data: dict, contenuto_ai: Optional[dict], output_dir: str) -> str:
    ai = contenuto_ai or {}
    doc = Document()
    _setup(doc)
    _footer(doc, "PIANO DI SICUREZZA E COORDINAMENTO (PSC) — Art. 100 D.Lgs. 81/2008 — All. XV")

    # ── Intestazione ──────────────────────────────────────────────────────────
    _header_bicolor(doc, "PIANO DI SICUREZZA E COORDINAMENTO — PSC")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PIANO DI SICUREZZA E COORDINAMENTO")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Art. 100 — D.Lgs. 9 aprile 2008, n. 81 — Allegato XV")
    r2.italic = True; r2.font.size = Pt(10); r2.font.color.rgb = GREY
    doc.add_paragraph()

    # ── Frontespizio ──────────────────────────────────────────────────────────
    front = doc.add_table(rows=7, cols=2)
    front.style = "Table Grid"
    front_data = [
        ("Opera",             v(form_data, "descrizione_opera")),
        ("Natura",            v(form_data, "natura_opera")),
        ("Indirizzo cantiere", f"{v(form_data,'indirizzo_cantiere')}, "
                               f"{v(form_data,'citta_cantiere')} ({v(form_data,'provincia_cantiere')})"),
        ("Committente",       f"{v(form_data,'committente_nome')} {v(form_data,'committente_cognome')} "
                               f"/ {v(form_data,'committente_ragione_sociale')}"),
        ("CSP",               f"{v(form_data,'csp_nome')} {v(form_data,'csp_cognome')}"),
        ("Importo lavori",    f"€ {v(form_data,'importo_lavori','—')}"),
        ("Costi sicurezza (non soggetti a ribasso)", f"€ {v(form_data,'costi_sicurezza','—')}"),
    ]
    for i, (lbl, val) in enumerate(front_data):
        _shd(front.cell(i, 0), BG_H)
        r = front.cell(i, 0).paragraphs[0].add_run(lbl)
        r.bold = True; r.font.size = Pt(9)
        front.cell(i, 1).paragraphs[0].add_run(val).font.size = Pt(9)
    doc.add_paragraph()

    # ── Indice ────────────────────────────────────────────────────────────────
    _section(doc, "INDICE")
    for item in [
        "1.  Identificazione e descrizione dell'opera",
        "2.  Individuazione dei soggetti con compiti di sicurezza",
        "3.  Analisi e valutazione dei rischi",
        "4.  Organizzazione del cantiere",
        "5.  Fasi lavorative e misure di prevenzione",
        "6.  Prescrizioni operative e DPI",
        "7.  Misure di coordinamento tra imprese",
        "8.  Gestione delle emergenze",
        "9.  Stima dei costi della sicurezza",
        "10. Firme",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ── 1. Identificazione opera ──────────────────────────────────────────────
    _section(doc, "IDENTIFICAZIONE E DESCRIZIONE DELL'OPERA", "1.")
    _note(doc, "Allegato XV, punto 2.1.2 lett. a — D.Lgs. 81/2008")
    _field(doc, "Natura dell'opera",   v(form_data, "natura_opera"))
    _field(doc, "Descrizione",         v(form_data, "descrizione_opera"))
    _field(doc, "Indirizzo cantiere",  v(form_data, "indirizzo_cantiere"))
    _field(doc, "Comune",
           f"{v(form_data,'citta_cantiere')} ({v(form_data,'provincia_cantiere')})")
    _field(doc, "Data inizio lavori",  v(form_data, "data_inizio"))
    _field(doc, "Data fine lavori",    v(form_data, "data_fine"))
    _field(doc, "Durata (uomini-giorno)", v(form_data, "uomini_giorno"))
    _field(doc, "N. max lavoratori contemporanei", v(form_data, "max_lavoratori"))
    _field(doc, "N. imprese esecutrici previste", v(form_data, "num_imprese"))
    _field(doc, "Importo lavori",      f"€ {v(form_data,'importo_lavori')}")

    # ── 2. Soggetti ───────────────────────────────────────────────────────────
    doc.add_page_break()
    _section(doc, "SOGGETTI CON COMPITI DI SICUREZZA", "2.")
    _note(doc, "Allegato XV, punto 2.1.2 lett. b — D.Lgs. 81/2008")

    _subsection(doc, "2.1  Committente")
    _field(doc, "Nome e Cognome",
           f"{v(form_data,'committente_nome')} {v(form_data,'committente_cognome')}")
    _field(doc, "Ragione Sociale",    v(form_data, "committente_ragione_sociale"))
    _field(doc, "Codice Fiscale",     v(form_data, "committente_cf"))
    _field(doc, "Indirizzo",          v(form_data, "committente_indirizzo"))
    _field(doc, "Telefono",           v(form_data, "committente_telefono"))
    _field(doc, "Email / PEC",        v(form_data, "committente_email"))

    _subsection(doc, "2.2  Responsabile dei Lavori")
    _field(doc, "Nome e Cognome", f"{v(form_data,'rl_nome','Coincide con committente')} {v(form_data,'rl_cognome','')}")
    _field(doc, "Qualifica",      v(form_data, "rl_qualifica"))
    _field(doc, "Telefono",       v(form_data, "rl_telefono"))

    _subsection(doc, "2.3  CSP — Coordinatore per la Progettazione")
    _field(doc, "Nome e Cognome",    f"{v(form_data,'csp_nome')} {v(form_data,'csp_cognome')}")
    _field(doc, "Ordine Prof.",      v(form_data, "csp_ordine"))
    _field(doc, "N. Iscrizione",     v(form_data, "csp_numero_ordine"))
    _field(doc, "Ultimo aggiornamento", v(form_data, "csp_data_aggiornamento"))
    _field(doc, "PEC",               v(form_data, "csp_pec"))

    _subsection(doc, "2.4  CSE — Coordinatore per l'Esecuzione")
    if v(form_data, "cse_nome"):
        _field(doc, "Nome e Cognome", f"{v(form_data,'cse_nome')} {v(form_data,'cse_cognome')}")
        _field(doc, "Ordine Prof.",   v(form_data, "cse_ordine"))
        _field(doc, "N. Iscrizione",  v(form_data, "cse_numero_ordine"))
        _field(doc, "PEC",            v(form_data, "cse_pec"))
    else:
        doc.add_paragraph("Da designare prima dell'inizio dei lavori").runs[0].italic = True

    _subsection(doc, "2.5  Imprese Esecutrici")
    imprese = form_data.get("imprese_esecutrici", [])
    if imprese:
        _table(doc,
            [[i.get("ragione_sociale",""), i.get("piva",""),
              f"{i.get('nome_dl','')} {i.get('cognome_dl','')}".strip(),
              f"{i.get('nome_rspp','')} {i.get('cognome_rspp','')}".strip(),
              i.get("attivita","")]
             for i in imprese],
            ["Ragione Sociale", "P.IVA", "Datore di Lavoro", "RSPP", "Attività"])
    else:
        doc.add_paragraph("[Elencare le imprese esecutrici previste]").runs[0].italic = True

    # ── 3. Analisi rischi ─────────────────────────────────────────────────────
    doc.add_page_break()
    _section(doc, "ANALISI E VALUTAZIONE DEI RISCHI", "3.")
    _note(doc, "Allegato XV, punto 2.1.2 lett. c — D.Lgs. 81/2008")

    for sub_title, ai_key, fallback in [
        ("3.1  Rischi derivanti dall'area di cantiere",
         "rischi_area",
         "Descrivere i rischi specifici dell'area: traffico, sottoservizi, "
         "condizioni del terreno, edifici adiacenti, linee aeree."),
        ("3.2  Rischi derivanti dalle lavorazioni",
         "rischi_lavorazioni",
         "Analizzare per ogni fase: caduta dall'alto, seppellimento, "
         "investimento, elettrocuzione, rumore, polveri, agenti chimici."),
        ("3.3  Rischi da interferenze tra imprese",
         "rischi_interferenze",
         "Identificare sovrapposizioni temporali e spaziali delle lavorazioni "
         "delle diverse imprese e le relative misure di gestione."),
    ]:
        _subsection(doc, sub_title)
        testo = ai.get(ai_key, fallback)
        p = doc.add_paragraph(testo)
        p.runs[0].font.size = Pt(10)

    # ── 4. Organizzazione cantiere ────────────────────────────────────────────
    doc.add_page_break()
    _section(doc, "ORGANIZZAZIONE DEL CANTIERE", "4.")
    _note(doc, "Allegato XV, punto 2.2.1 — D.Lgs. 81/2008")

    org_items = [
        ("4.1  Viabilità e accessi",          "viabilita"),
        ("4.2  Zone di carico e scarico",      "zone_carico"),
        ("4.3  Zone di stoccaggio materiali",  "zone_stoccaggio"),
        ("4.4  Recinzione e segnaletica",      "recinzione"),
        ("4.5  Impianti di cantiere",          "impianti"),
        ("4.6  Servizi igienico-assistenziali","servizi_igienici"),
        ("4.7  Primo soccorso e antincendio",  "primo_soccorso"),
    ]
    for sub_title, ai_key in org_items:
        _subsection(doc, sub_title)
        testo = ai.get(ai_key, f"[Descrivere: {sub_title}]")
        p = doc.add_paragraph(testo)
        p.runs[0].font.size = Pt(10)

    # ── 5. Fasi lavorative ────────────────────────────────────────────────────
    doc.add_page_break()
    _section(doc, "FASI LAVORATIVE E MISURE DI PREVENZIONE", "5.")
    _note(doc, "Allegato XV, punto 2.2.2 — D.Lgs. 81/2008")

    fasi_default = [
        {"nome": "Allestimento cantiere",     "rischi": "Investimento, caduta materiali", "misure": "Recinzione, segnaletica, DPI base", "dpi": "Casco, guanti, scarpe S3"},
        {"nome": "Scavi e fondazioni",        "rischi": "Seppellimento, franamento pareti, investimento mezzi", "misure": "Sbadacchiatura scavi >1.5m, distanze di sicurezza, segnalazione bordi", "dpi": "Casco, scarpe S3, gilet av"},
        {"nome": "Strutture in elevazione",   "rischi": "Caduta dall'alto, caduta materiali", "misure": "Ponteggi omologati, reti, parapetti, ordini di getto", "dpi": "Casco, imbragatura, scarpe S3"},
        {"nome": "Murature e tamponamenti",   "rischi": "Caduta dall'alto, movimentazione manuale carichi", "misure": "Trabattelli, limiti peso, ausili meccanici", "dpi": "Casco, guanti, scarpe S3"},
        {"nome": "Impianti",                  "rischi": "Elettrocuzione, caduta dall'alto", "misure": "Quadri CEI, distanze linee, lavori fuori tensione", "dpi": "Guanti diel., scarpe isol."},
        {"nome": "Finiture",                  "rischi": "Caduta dall'alto, agenti chimici (vernici, colle)", "misure": "Trabattelli, ventilazione, SDS prodotti", "dpi": "Casco, maschere FFP, guanti"},
        {"nome": "Smobilizzo cantiere",       "rischi": "Caduta materiali, investimento", "misure": "Rimozione progressiva, viabilità controllata", "dpi": "Casco, guanti, scarpe S3"},
    ]

    fasi = form_data.get("fasi_lavorative", fasi_default)
    for i, fase in enumerate(fasi, 1):
        _subsection(doc, f"Fase {i}: {fase.get('nome','')}")
        _table(doc, [
            ["Rischi principali",        fase.get("rischi","—")],
            ["Misure di prevenzione",    fase.get("misure","—")],
            ["DPI richiesti",            fase.get("dpi","—")],
            ["Imprese coinvolte",        fase.get("imprese","Da definire")],
            ["Durata stimata",           fase.get("durata","—")],
        ], ["Voce", "Contenuto"])

    # ── 6. DPI ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    _section(doc, "PRESCRIZIONI OPERATIVE E DPI", "6.")
    _note(doc, "Allegato XV, punto 2.1.2 lett. d — Allegato VIII — D.Lgs. 81/2008")
    _table(doc, [
        ["Testa",        "Casco di protezione",         "EN 397",          "Sempre in cantiere"],
        ["Piedi",        "Scarpe antinf. cat. S3",      "EN ISO 20345",    "Sempre in cantiere"],
        ["Mani",         "Guanti di protezione",         "EN 420",          "Durante lavorazioni"],
        ["Occhi",        "Occhiali / Visiera",           "EN 166",          "Taglio, molatura, saldatura"],
        ["Udito",        "Cuffie / Tappi auricolari",   "EN 352",          "Se Lex,8h > 80 dB(A)"],
        ["Vie resp.",    "Maschera FFP2 / FFP3",        "EN 149",          "Polveri, vernici, agenti chimici"],
        ["Corpo",        "Gilet alta visibilità cl.2",  "EN ISO 20471",    "Lavori stradali"],
        ["Anticaduta",   "Imbragatura + linea vita",    "EN 361 / EN 354", "Lavori in quota > 2m"],
    ], ["Parte corpo protetta", "DPI", "Norma", "Quando"])

    # ── 7. Coordinamento ──────────────────────────────────────────────────────
    _section(doc, "MISURE DI COORDINAMENTO TRA IMPRESE", "7.")
    _note(doc, "Allegato XV, punto 2.3 — D.Lgs. 81/2008")
    testo = ai.get("coordinamento",
        "Le imprese esecutrici sono tenute a coordinarsi tra loro per:\n"
        "• Utilizzo comune di apprestamenti (ponteggi, trabattelli, protezioni collettive)\n"
        "• Accesso e viabilità interna al cantiere\n"
        "• Gestione aree di deposito e stoccaggio\n"
        "• Modalità e orari di lavoro per evitare interferenze pericolose\n"
        "• Comunicazione immediata al CSE di qualsiasi variazione delle lavorazioni.\n\n"
        "Il CSE convoca riunioni di coordinamento con cadenza almeno mensile "
        "o al verificarsi di variazioni significative (art. 92 D.Lgs. 81/2008).")
    doc.add_paragraph(testo).runs[0].font.size = Pt(10)

    # ── 8. Emergenze ──────────────────────────────────────────────────────────
    _section(doc, "GESTIONE DELLE EMERGENZE", "8.")
    _note(doc, "D.M. 388/2003 — D.M. 10/03/1998 — D.Lgs. 81/2008")
    _table(doc, [
        ["Numero Unico Emergenze", "112"],
        ["SUEM / Pronto Soccorso",  "118"],
        ["Vigili del Fuoco",        "115"],
        ["INAIL (infortuni)",       "803.164"],
        ["ASL competente",         v(form_data, "asl_destinataria", "—")],
        ["Ospedale/PS più vicino", v(form_data, "ospedale_vicino", "—")],
        ["Addetti primo soccorso", v(form_data, "addetti_ps", "Da nominare")],
        ["Addetti antincendio",    v(form_data, "addetti_ai", "Da nominare")],
    ], ["Voce", "Riferimento / Nominativo"])
    testo_em = ai.get("emergenze_procedure",
        "In caso di infortunio: prestare primo soccorso, chiamare il 118, mettere in sicurezza "
        "l'area, avvisare immediatamente il CSE e il Datore di Lavoro, compilare la denuncia "
        "di infortunio INAIL entro 24h per eventi > 3 giorni. "
        "Presidi: cassetta PS (D.M. 388/2003), estintori polvere ABC 6kg ogni 200mq.")
    doc.add_paragraph(testo_em).runs[0].font.size = Pt(10)

    # ── 9. Costi sicurezza ────────────────────────────────────────────────────
    doc.add_page_break()
    _section(doc, "STIMA DEI COSTI DELLA SICUREZZA", "9.")
    _note(doc, "Allegato XV, punto 4 — D.Lgs. 81/2008")
    p_avv = doc.add_paragraph(
        "I costi della sicurezza sono ESCLUSI dal ribasso d'asta — art. 100 co. 1, D.Lgs. 81/2008"
    )
    p_avv.runs[0].bold = True; p_avv.runs[0].font.color.rgb = RED

    _table(doc, [
        ["Apprestamenti (ponteggi, trabattelli, etc.)",            f"€ {v(form_data,'costo_apprestamenti','—')}"],
        ["Misure preventive, protettive e DPI aggiuntivi",         f"€ {v(form_data,'costo_dpi','—')}"],
        ["Impianto di terra e protezione scariche atmosferiche",   f"€ {v(form_data,'costo_impianti','—')}"],
        ["Mezzi e servizi di protezione collettiva",               f"€ {v(form_data,'costo_protezioni','—')}"],
        ["Procedure specifiche per imprese speciali",              f"€ {v(form_data,'costo_procedure','—')}"],
        ["Presidi di primo soccorso e antincendio",                f"€ {v(form_data,'costo_ps_ai','—')}"],
        ["Cartellonistica di sicurezza",                           f"€ {v(form_data,'costo_cartelli','—')}"],
        ["TOTALE COSTI SICUREZZA",                                 f"€ {v(form_data,'costi_sicurezza','—')}"],
    ], ["Voce di costo", "Importo"])

    # ── 10. Firme ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    _section(doc, "FIRME", "10.")
    p_decl = doc.add_paragraph(
        "Il sottoscritto CSP dichiara che il presente PSC è stato redatto in conformità "
        "all'art. 100 e all'Allegato XV del D.Lgs. 81/2008, con il coinvolgimento del "
        "Responsabile dei Lavori e previa presa visione dell'area di cantiere."
    )
    p_decl.runs[0].font.size = Pt(10)
    doc.add_paragraph()
    _table(doc, [
        ["Il Committente / Responsabile dei Lavori", f"{v(form_data,'committente_nome')} {v(form_data,'committente_cognome')}", "_____________________________"],
        ["Il Coordinatore per la Progettazione (CSP)", f"{v(form_data,'csp_nome')} {v(form_data,'csp_cognome')}", "_____________________________"],
        ["Luogo e data", f"{v(form_data,'citta_cantiere')},  _____ / _____ / _______", ""],
    ], ["Qualifica", "Nome e Cognome", "Firma"])
    doc.add_paragraph()
    _warning(doc)

    return _save(doc, "PSC", output_dir, v(form_data, "citta_cantiere"))


# ═══════════════════════════════════════════════════════════════════════════════
# POS — PIANO OPERATIVO DI SICUREZZA
# ═══════════════════════════════════════════════════════════════════════════════

def genera_pos(form_data: dict, contenuto_ai: Optional[dict], output_dir: str) -> str:
    ai = contenuto_ai or {}
    doc = Document()
    _setup(doc)
    _footer(doc, "PIANO OPERATIVO DI SICUREZZA (POS) — Art. 101 D.Lgs. 81/2008 — All. XV punto 3")

    # ── Intestazione ──────────────────────────────────────────────────────────
    _header_bicolor(doc, "PIANO OPERATIVO DI SICUREZZA — POS")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PIANO OPERATIVO DI SICUREZZA")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Art. 101 — D.Lgs. 9 aprile 2008, n. 81 — Allegato XV punto 3")
    r2.italic = True; r2.font.size = Pt(10); r2.font.color.rgb = GREY
    doc.add_paragraph()

    # ── Frontespizio ──────────────────────────────────────────────────────────
    front = doc.add_table(rows=5, cols=2)
    front.style = "Table Grid"
    for i, (lbl, val) in enumerate([
        ("Impresa Esecutrice",  v(form_data, "impresa_ragione_sociale")),
        ("P.IVA",               v(form_data, "impresa_piva")),
        ("Cantiere",            f"{v(form_data,'indirizzo_cantiere')}, {v(form_data,'citta_cantiere')}"),
        ("Datore di Lavoro",    f"{v(form_data,'nome_dl')} {v(form_data,'cognome_dl')}"),
        ("Data redazione",      v(form_data, "data_redazione", datetime.now().strftime("%d/%m/%Y"))),
    ]):
        _shd(front.cell(i, 0), BG_H)
        r = front.cell(i, 0).paragraphs[0].add_run(lbl)
        r.bold = True; r.font.size = Pt(9)
        front.cell(i, 1).paragraphs[0].add_run(val).font.size = Pt(9)
    doc.add_paragraph()

    # ── 1. Dati impresa ───────────────────────────────────────────────────────
    _section(doc, "DATI IDENTIFICATIVI DELL'IMPRESA ESECUTRICE", "1.")
    _note(doc, "Allegato XV, punto 3.2.1 lett. a — D.Lgs. 81/2008")
    _field(doc, "Ragione Sociale",  v(form_data, "impresa_ragione_sociale"))
    _field(doc, "Codice Fiscale",   v(form_data, "impresa_cf"))
    _field(doc, "P.IVA",            v(form_data, "impresa_piva"))
    _field(doc, "Sede legale",      v(form_data, "impresa_indirizzo"))
    _field(doc, "Comune",
           f"{v(form_data,'impresa_citta')} ({v(form_data,'impresa_provincia')}) "
           f"{v(form_data,'impresa_cap')}")
    _field(doc, "Telefono",         v(form_data, "impresa_telefono"))
    _field(doc, "Email / PEC",      v(form_data, "impresa_email"))
    _field(doc, "Iscr. CCIAA",      f"{v(form_data,'impresa_cciaa')} n. {v(form_data,'impresa_numero_cciaa')}")
    _field(doc, "Posizione INAIL (PAT)", v(form_data, "impresa_inail_pat"))
    _field(doc, "Cassa Edile / INPS",   v(form_data, "impresa_cassa_edile"))
    _field(doc, "CCNL applicato",   v(form_data, "impresa_ccnl", "CCNL Edilizia Industria"))

    # ── 2. Figure sicurezza ───────────────────────────────────────────────────
    _section(doc, "FIGURE DELLA SICUREZZA", "2.")
    _note(doc, "Allegato XV, punto 3.2.1 lett. d — D.Lgs. 81/2008")
    _table(doc, [
        ["Datore di Lavoro",      f"{v(form_data,'nome_dl')} {v(form_data,'cognome_dl')}",         "Art. 2 co.1 lett. b"],
        ["RSPP",                  f"{v(form_data,'nome_rspp')} {v(form_data,'cognome_rspp')}",     "Art. 17, 31"],
        ["Medico Competente",     f"{v(form_data,'nome_mc')} {v(form_data,'cognome_mc')}",         "Art. 38"],
        ["RLS / RLST",            f"{v(form_data,'nome_rls')} {v(form_data,'cognome_rls')}",       "Art. 47-50"],
        ["Preposto cantiere",     f"{v(form_data,'nome_preposto','—')} {v(form_data,'cognome_preposto','')}", "Art. 19"],
        ["Addetto Primo Soccorso", v(form_data, "addetto_ps", "Da nominare"),                      "D.M. 388/2003"],
        ["Addetto Antincendio",   v(form_data, "addetto_ai", "Da nominare"),                       "D.M. 10/03/1998"],
    ], ["Ruolo", "Nome e Cognome", "Rif. normativo"])

    # ── 3. Attività nel cantiere ──────────────────────────────────────────────
    _section(doc, "DESCRIZIONE DELLE ATTIVITÀ NEL CANTIERE", "3.")
    _note(doc, "Allegato XV, punto 3.2.1 lett. b e c — D.Lgs. 81/2008")
    _field(doc, "Cantiere",
           f"{v(form_data,'indirizzo_cantiere')}, {v(form_data,'citta_cantiere')}")
    _field(doc, "Attività svolta dall'impresa",  v(form_data, "attivita_cantiere"))
    _field(doc, "Fasi di lavoro proprie",        v(form_data, "fasi_proprie"))
    _field(doc, "Periodo di intervento",         v(form_data, "periodo_intervento"))
    _field(doc, "N. lavoratori impiegati",       v(form_data, "num_lavoratori"))

    # ── 4. Elenco lavoratori ──────────────────────────────────────────────────
    doc.add_page_break()
    _section(doc, "ELENCO DEI LAVORATORI", "4.")
    _note(doc, "Allegato XV, punto 3.2.1 lett. e — Art. 90 co.9 lett. c — D.Lgs. 81/2008")
    lavoratori = form_data.get("lavoratori", [])
    if lavoratori:
        _table(doc,
            [[l.get("nome",""), l.get("cognome",""), l.get("cf",""),
              l.get("mansione",""), l.get("idoneita",""), l.get("formazione","")]
             for l in lavoratori],
            ["Nome", "Cognome", "Cod. Fiscale", "Mansione", "Idoneità sanitaria", "Formazione"])
    else:
        p = doc.add_paragraph(
            "[Inserire l'elenco nominativo dei lavoratori con mansione, "
            "idoneità sanitaria e attestati di formazione]"
        )
        p.runs[0].italic = True; p.runs[0].font.color.rgb = RGBColor(0xCC, 0x44, 0x00)

    # ── 5. Macchine e attrezzature ────────────────────────────────────────────
    _section(doc, "MACCHINE, ATTREZZATURE E IMPIANTI", "5.")
    _note(doc, "Allegato XV, punto 3.2.1 lett. f — D.Lgs. 81/2008")
    macchine = form_data.get("macchine", [])
    if macchine:
        _table(doc,
            [[m.get("nome",""), m.get("marca",""), m.get("anno",""),
              m.get("matricola",""), m.get("operatore","")]
             for m in macchine],
            ["Attrezzatura", "Marca/Modello", "Anno", "N. Matricola", "Operatore abilitato"])
    else:
        _table(doc, [
            ["Betoniera",              "—", "—", "—", "—"],
            ["Sega circolare",         "—", "—", "—", "—"],
            ["Smerigliatrice angolare","—", "—", "—", "—"],
            ["Trapano a percussione",  "—", "—", "—", "—"],
        ], ["Attrezzatura", "Marca/Modello", "Anno", "Matricola", "Operatore"])

    # ── 6. Sostanze pericolose ────────────────────────────────────────────────
    _section(doc, "SOSTANZE E PREPARATI PERICOLOSI", "6.")
    _note(doc, "Allegato XV, punto 3.2.1 lett. g — Reg. REACH (CE) 1907/2006 — CLP 1272/2008")
    p = doc.add_paragraph(
        "Per ogni sostanza pericolosa utilizzata è disponibile la Scheda Dati di Sicurezza "
        "(SDS) conforme al Regolamento (CE) n. 1907/2006 (REACH)."
    )
    p.runs[0].font.size = Pt(10)
    _table(doc, [
        ["Cemento e calce",         "Corrosivo (skin/occhi)",  "SDS allegata", "Guanti, occhiali, maschera P2"],
        ["Primer / sigillanti",     "Irritante H315/H319",     "SDS allegata", "Guanti, occhiali, ventilazione"],
        ["Solventi (se utilizzati)","Infiammabile H225",       "SDS allegata", "Guanti resist., ventilazione"],
    ], ["Sostanza/Preparato", "Classificazione pericolo", "Documentazione", "Misure di protezione"])

    # ── 7. Valutazione rischi ─────────────────────────────────────────────────
    doc.add_page_break()
    _section(doc, "VALUTAZIONE DEI RISCHI SPECIFICI", "7.")
    _note(doc, "Allegato XV, punto 3.2.1 lett. h — D.Lgs. 81/2008")
    testo_rischi = ai.get("rischi_specifici", "")
    if testo_rischi:
        doc.add_paragraph(testo_rischi).runs[0].font.size = Pt(10)
    else:
        _table(doc, [
            ["Caduta dall'alto",              "ALTO",  "Parapetti, imbragature, linea vita",          "MEDIO"],
            ["Investimento mezzi meccanici",  "ALTO",  "Viabilità separata, segnaletica, barriere",   "BASSO"],
            ["Elettrocuzione",                "MEDIO", "Quadri CEI 60439, lavori fuori tensione",     "BASSO"],
            ["Movim. manuale carichi",        "MEDIO", "Ausili meccanici, formazione movimentazione", "BASSO"],
            ["Rumore",                        "MEDIO", "Otoprotettori, rotazione turni di lavoro",    "BASSO"],
            ["Polveri / agenti chimici",      "MEDIO", "Maschere FFP2/3, ventilazione, SDS",          "BASSO"],
            ["Scivolamento caduta stessa quota","BASSO","Pulizia vie transito, calzature S3",         "BASSO"],
        ], ["Rischio", "Livello iniziale", "Misure adottate", "Livello residuo"])

    # ── 8. DPI ────────────────────────────────────────────────────────────────
    _section(doc, "DISPOSITIVI DI PROTEZIONE INDIVIDUALE (DPI)", "8.")
    _note(doc, "Allegato XV, punto 3.2.1 lett. i — Allegato VIII — D.Lgs. 81/2008")
    _table(doc, [
        ["Casco (EN 397)",              "Sempre in cantiere",                    "Tutti i lavoratori"],
        ["Scarpe S3 (EN ISO 20345)",    "Sempre in cantiere",                    "Tutti i lavoratori"],
        ["Guanti (EN 420)",             "Durante tutte le lavorazioni",          "Tutti i lavoratori"],
        ["Occhiali (EN 166)",           "Taglio, molatura, saldatura",           "Addetti"],
        ["Cuffie/tappi (EN 352)",       "Se Lex,8h > 80 dB(A)",                 "Addetti macchine rumorose"],
        ["Maschera FFP2 (EN 149)",      "Polveri, vernici, agenti chimici",      "Addetti"],
        ["Imbragatura (EN 361)",        "Lavori in quota > 2m",                  "Addetti lavori in quota"],
    ], ["DPI", "Quando utilizzarlo", "Chi lo utilizza"])

    # ── 9. Formazione ─────────────────────────────────────────────────────────
    _section(doc, "FORMAZIONE E ADDESTRAMENTO", "9.")
    _note(doc, "Art. 37 — Accordo Stato-Regioni 21/12/2011 — D.Lgs. 81/2008")
    p = doc.add_paragraph(
        "Tutti i lavoratori sono in possesso dell'attestato di formazione generale (4 ore) "
        "e specifica per il settore edile ad alto rischio (12 ore) conformemente all'Accordo "
        "Stato-Regioni del 21/12/2011. I documenti originali sono disponibili in cantiere."
    )
    p.runs[0].font.size = Pt(10)

    # ── 10. Emergenze ─────────────────────────────────────────────────────────
    _section(doc, "GESTIONE DELLE EMERGENZE", "10.")
    _note(doc, "D.M. 388/2003 — D.M. 10/03/1998")
    _table(doc, [
        ["N. unico emergenze",    "112"],
        ["SUEM 118 / PS",         "118"],
        ["Vigili del Fuoco",      "115"],
        ["Ospedale più vicino",   v(form_data, "ospedale_vicino", "—")],
        ["Addetto PS",            v(form_data, "addetto_ps", "Da nominare")],
        ["Addetto antincendio",   v(form_data, "addetto_ai", "Da nominare")],
    ], ["Voce", "Riferimento / Nominativo"])
    testo_em = ai.get("gestione_emergenze",
        "In caso di infortunio: applicare primo soccorso, chiamare il 118, "
        "mettere in sicurezza l'area, avvisare il DL e il CSE. "
        "In caso di incendio: azionare l'allarme, evacuare il cantiere, chiamare il 115.")
    doc.add_paragraph(testo_em).runs[0].font.size = Pt(10)

    # ── 11. Procedure operative ───────────────────────────────────────────────
    _section(doc, "PROCEDURE DI SICUREZZA SPECIFICHE", "11.")
    _note(doc, "Allegato XV, punto 3.2.1 lett. l — D.Lgs. 81/2008")
    testo_proc = ai.get("procedure_operative",
        "[Il professionista deve integrare con le procedure specifiche adottate dall'impresa "
        "per le proprie lavorazioni: uso ponteggi, movimentazione carichi, uso macchine, "
        "lavorazioni in quota, scavi, ecc.]")
    p = doc.add_paragraph(testo_proc)
    p.runs[0].font.size = Pt(10)
    if "professionista" in testo_proc:
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0x44, 0x00)

    # ── 12. Firme ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    _section(doc, "FIRME E DICHIARAZIONI", "12.")
    p_decl = doc.add_paragraph(
        "Il sottoscritto Datore di Lavoro dichiara che il presente POS è stato redatto "
        "in conformità all'art. 101 e all'Allegato XV del D.Lgs. 81/2008, con il "
        "coinvolgimento del Servizio di Prevenzione e Protezione e previa consultazione "
        "del RLS/RLST ai sensi dell'art. 102."
    )
    p_decl.runs[0].font.size = Pt(10)
    doc.add_paragraph()
    _table(doc, [
        ["Datore di Lavoro",         f"{v(form_data,'nome_dl')} {v(form_data,'cognome_dl')}",     "_____________________________"],
        ["RSPP",                     f"{v(form_data,'nome_rspp')} {v(form_data,'cognome_rspp')}", "_____________________________"],
        ["RLS/RLST (consultato)",    f"{v(form_data,'nome_rls')} {v(form_data,'cognome_rls')}",   "_____________________________"],
        ["Luogo e data",             f"{v(form_data,'citta_cantiere')},  _____ / _____ / _______",""],
    ], ["Qualifica", "Nome e Cognome", "Firma"])
    doc.add_paragraph()
    _warning(doc)

    return _save(doc, "POS",
                 output_dir,
                 v(form_data, "impresa_ragione_sociale", "impresa")[:25])

from fastapi import APIRouter, UploadFile, File, HTTPException
from typing import List
import anthropic
import base64
import json
import io

router = APIRouter(prefix="/api/estrazione", tags=["estrazione"])

SYSTEM_PROMPT = """Sei un assistente specializzato nell'estrazione di dati da documenti edilizi e professionali italiani.
Leggi attentamente il documento e estrai tutti i dati strutturati presenti.
Rispondi SEMPRE e SOLO con un JSON valido, senza markdown, senza testo aggiuntivo prima o dopo."""

EXTRACTION_PROMPT = """Analizza questo documento ed estrai tutti i dati presenti relativi a un cantiere edile italiano.

Cerca e estrai i seguenti dati (includi SOLO quelli effettivamente presenti, non inventare mai):

COMMITTENTE: nome, cognome, ragione_sociale, codice_fiscale, piva, indirizzo, citta, provincia, cap, telefono, email, pec, tipo (persona_fisica o persona_giuridica)

IMPRESA ESECUTRICE: ragione_sociale, piva, codice_fiscale, indirizzo, citta, provincia, telefono, email, cciaa, numero_cciaa, inail_pat, cassa_edile, ccnl, nome_dl, cognome_dl, nome_rspp, cognome_rspp, nome_mc, cognome_mc, nome_rls, cognome_rls

CANTIERE: indirizzo_cantiere, citta_cantiere, provincia_cantiere, cap_cantiere

OPERA: natura_opera, descrizione_opera, destinazione_uso

LAVORI: data_inizio, data_fine, durata_lavori, importo_lavori, importo_sicurezza, num_lavoratori, uomini_giorno, max_lavoratori, fasi_descrizione, lavorazioni_critiche, asl_destinataria

COORDINATORI CSP: csp_nome, csp_cognome, csp_ordine, csp_numero_ordine, csp_pec, csp_telefono
COORDINATORI CSE: cse_nome, cse_cognome, cse_ordine, cse_numero_ordine, cse_pec, cse_telefono

Rispondi ESCLUSIVAMENTE con questo JSON (ometti i campi non trovati):
{
  "fonte": "descrizione breve del tipo di documento rilevato (es: Contratto d appalto, Visura camerale, Documento identita, Progetto architettonico, Computo metrico)",
  "committente": {},
  "impresa": {},
  "cantiere": {},
  "opera": {},
  "lavori": {},
  "coordinatori": {},
  "note": "eventuali osservazioni importanti"
}"""


def estrai_testo_docx(contenuto: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(contenuto))
    testo = []
    for para in doc.paragraphs:
        if para.text.strip():
            testo.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            riga = []
            for cell in row.cells:
                if cell.text.strip():
                    riga.append(cell.text.strip())
            if riga:
                testo.append(' | '.join(riga))
    return '\n'.join(testo)


def estrai_testo_xlsx(contenuto: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(contenuto), data_only=True)
    testo = []
    for sheet in wb.worksheets:
        testo.append(f"[Foglio: {sheet.title}]")
        for row in sheet.iter_rows():
            riga = [str(cell.value) for cell in row if cell.value is not None]
            if riga:
                testo.append(' | '.join(riga))
    return '\n'.join(testo)


def parse_json_risposta(testo: str) -> dict:
    testo = testo.strip()
    if testo.startswith('```'):
        lines = testo.split('\n')
        testo = '\n'.join(lines[1:])
        testo = testo.rsplit('```', 1)[0]
    return json.loads(testo.strip())


async def analizza_singolo_documento(file: UploadFile) -> dict:
    client = anthropic.Anthropic()
    contenuto = await file.read()
    nome = file.filename.lower()

    try:
        if nome.endswith('.pdf'):
            b64 = base64.standard_b64encode(contenuto).decode('utf-8')
            risposta = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=2000,
                system=SYSTEM_PROMPT,
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "document",
                            "source": {
                                "type": "base64",
                                "media_type": "application/pdf",
                                "data": b64
                            }
                        },
                        {"type": "text", "text": EXTRACTION_PROMPT}
                    ]
                }]
            )

        elif nome.endswith(('.jpg', '.jpeg', '.png')):
            media_type = "image/png" if nome.endswith('.png') else "image/jpeg"
            b64 = base64.standard_b64encode(contenuto).decode('utf-8')
            risposta = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=2000,
                system=SYSTEM_PROMPT,
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": b64
                            }
                        },
                        {"type": "text", "text": EXTRACTION_PROMPT}
                    ]
                }]
            )

        elif nome.endswith('.docx'):
            testo = estrai_testo_docx(contenuto)
            risposta = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=2000,
                system=SYSTEM_PROMPT,
                messages=[{
                    "role": "user",
                    "content": f"Documento Word estratto:\n\n{testo}\n\n{EXTRACTION_PROMPT}"
                }]
            )

        elif nome.endswith(('.xlsx', '.xls')):
            testo = estrai_testo_xlsx(contenuto)
            risposta = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=2000,
                system=SYSTEM_PROMPT,
                messages=[{
                    "role": "user",
                    "content": f"Foglio Excel estratto:\n\n{testo}\n\n{EXTRACTION_PROMPT}"
                }]
            )

        else:
            return {
                "_filename": file.filename,
                "_errore": f"Formato non supportato: {file.filename}"
            }

        dati = parse_json_risposta(risposta.content[0].text)
        dati["_filename"] = file.filename
        return dati

    except Exception as e:
        return {
            "_filename": file.filename,
            "_errore": str(e)
        }


def unisci_risultati(risultati: list) -> dict:
    """
    Unisce i dati estratti da più documenti.
    Rileva conflitti quando lo stesso campo ha valori diversi da fonti diverse.
    """
    categorie = ['committente', 'impresa', 'cantiere', 'opera', 'lavori', 'coordinatori']
    merged = {cat: {} for cat in categorie}
    conflitti = []
    documenti_ok = []
    errori = []

    for r in risultati:
        if '_errore' in r:
            errori.append({"file": r.get('_filename', '?'), "errore": r['_errore']})
            continue

        filename = r.get('_filename', r.get('fonte', 'sconosciuto'))
        documenti_ok.append(filename)

        for cat in categorie:
            if cat not in r or not isinstance(r[cat], dict):
                continue
            for chiave, valore in r[cat].items():
                if valore is None:
                    continue
                valore_str = str(valore).strip()
                if not valore_str or valore_str.lower() in ('none', 'null', ''):
                    continue

                if chiave not in merged[cat]:
                    merged[cat][chiave] = {
                        "valore": valore_str,
                        "fonte": filename,
                        "conflitti": []
                    }
                else:
                    esistente = merged[cat][chiave]["valore"]
                    if valore_str.lower() != esistente.lower():
                        conflitti.append({
                            "campo": f"{cat}.{chiave}",
                            "valore_a": esistente,
                            "fonte_a": merged[cat][chiave]["fonte"],
                            "valore_b": valore_str,
                            "fonte_b": filename
                        })
                        merged[cat][chiave]["conflitti"].append({
                            "valore": valore_str,
                            "fonte": filename
                        })

    totale_dati = sum(len(campi) for campi in merged.values())

    return {
        "dati": merged,
        "conflitti": conflitti,
        "documenti_analizzati": documenti_ok,
        "errori": errori,
        "riepilogo": {
            "dati_estratti": totale_dati,
            "conflitti": len(conflitti),
            "documenti_analizzati": len(documenti_ok),
            "documenti_con_errori": len(errori)
        }
    }


@router.post("/analizza")
async def analizza_documenti(files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(status_code=400, detail="Nessun file caricato")

    risultati = []
    for file in files:
        r = await analizza_singolo_documento(file)
        risultati.append(r)

    return unisci_risultati(risultati)